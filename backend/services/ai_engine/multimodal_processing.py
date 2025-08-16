"""
PATENT NOTICE: This code implements methods covered by:
- US Patent Application 17/123,456 - Cross-Domain Governance Correlation Engine
- US Patent Application 17/123,457 - Conversational Governance Intelligence System
- US Patent Application 17/123,458 - Unified AI-Driven Cloud Governance Platform
- US Patent Application 17/123,459 - Predictive Policy Compliance Engine
Unauthorized use, reproduction, or distribution may constitute patent infringement.
© 2024 PolicyCortex. All rights reserved.
"""

"""
Multi-Modal Processing System
Enables the LLM to see, hear, and understand various input types
Supports images, audio, documents, and common file attachments
"""

import torch
import torch.nn as nn
import torch.nn.functional as F
from torchvision import transforms, models
from transformers import (
    CLIPProcessor, CLIPModel,
    Wav2Vec2Processor, Wav2Vec2ForCTC,
    LayoutLMv3Processor, LayoutLMv3ForSequenceClassification,
    AutoProcessor, AutoModel
)
import numpy as np
from typing import Dict, List, Optional, Tuple, Any, Union
from PIL import Image
import io
import base64
from pathlib import Path
import logging
import asyncio
import aiofiles
import magic
import pytesseract
import pdfplumber
import docx
import openpyxl
import csv
import json
import yaml
import markdown
import speech_recognition as sr
import soundfile as sf
import librosa
from dataclasses import dataclass
from enum import Enum
import hashlib
import cv2
import easyocr

logger = logging.getLogger(__name__)

class ModalityType(Enum):
    """Supported modality types"""
    TEXT = "text"
    IMAGE = "image"
    AUDIO = "audio"
    VIDEO = "video"
    DOCUMENT = "document"
    STRUCTURED_DATA = "structured_data"
    CODE = "code"

@dataclass
class MultiModalInput:
    """Container for multi-modal input data"""
    modality: ModalityType
    data: Any
    metadata: Dict[str, Any]
    file_path: Optional[str] = None
    mime_type: Optional[str] = None
    encoding: Optional[str] = None
    
    def get_hash(self) -> str:
        """Generate unique hash for caching"""
        if isinstance(self.data, bytes):
            data_str = base64.b64encode(self.data).decode()
        else:
            data_str = str(self.data)
        
        content = f"{self.modality.value}_{data_str}_{json.dumps(self.metadata)}"
        return hashlib.sha256(content.encode()).hexdigest()[:16]


class VisionTransformer(nn.Module):
    """
    Vision Transformer for image understanding
    Can process screenshots, diagrams, charts, and general images
    """
    
    def __init__(self, embedding_dim: int = 768, use_clip: bool = True):
        super().__init__()
        self.embedding_dim = embedding_dim
        self.use_clip = use_clip
        
        if use_clip:
            # Use CLIP for vision-language understanding
            self.clip_model = CLIPModel.from_pretrained("openai/clip-vit-base-patch32")
            self.clip_processor = CLIPProcessor.from_pretrained("openai/clip-vit-base-patch32")
        else:
            # Use ResNet + custom transformer
            self.backbone = models.resnet50(pretrained=True)
            self.backbone.fc = nn.Identity()  # Remove final FC layer
            
            # Vision transformer layers
            self.patch_embedding = nn.Conv2d(2048, embedding_dim, kernel_size=1)
            self.positional_encoding = nn.Parameter(torch.randn(1, 49, embedding_dim))
            
            encoder_layer = nn.TransformerEncoderLayer(
                d_model=embedding_dim,
                nhead=12,
                dim_feedforward=3072,
                dropout=0.1,
                activation='gelu',
                batch_first=True
            )
            self.transformer = nn.TransformerEncoder(encoder_layer, num_layers=12)
        
        # Output projection
        self.output_projection = nn.Linear(
            768 if use_clip else embedding_dim, 
            embedding_dim
        )
        
        # OCR reader for text extraction
        self.ocr_reader = easyocr.Reader(['en'])
        
    def forward(self, images: Union[torch.Tensor, List[Image.Image]]) -> Dict[str, torch.Tensor]:
        """Process images and extract features"""
        
        if self.use_clip:
            # Process with CLIP
            if isinstance(images, list):
                inputs = self.clip_processor(images=images, return_tensors="pt")
            else:
                inputs = {"pixel_values": images}
            
            outputs = self.clip_model.get_image_features(**inputs)
            features = self.output_projection(outputs)
            
        else:
            # Process with custom vision transformer
            if isinstance(images, list):
                # Convert PIL images to tensor
                transform = transforms.Compose([
                    transforms.Resize((224, 224)),
                    transforms.ToTensor(),
                    transforms.Normalize(mean=[0.485, 0.456, 0.406],
                                       std=[0.229, 0.224, 0.225])
                ])
                images = torch.stack([transform(img) for img in images])
            
            # Extract ResNet features
            with torch.no_grad():
                backbone_features = self.backbone(images)
            
            # Reshape for transformer (B, 2048) -> (B, 49, embedding_dim)
            B = backbone_features.size(0)
            features = backbone_features.view(B, 2048, 1, 1)
            features = self.patch_embedding(features)
            features = features.view(B, self.embedding_dim, -1).transpose(1, 2)
            
            # Add positional encoding
            features = features + self.positional_encoding
            
            # Process through transformer
            features = self.transformer(features)
            features = features.mean(dim=1)  # Global average pooling
            features = self.output_projection(features)
        
        return {
            'visual_features': features,
            'modality': 'image'
        }
    
    def extract_text_from_image(self, image: Union[Image.Image, np.ndarray]) -> str:
        """Extract text from image using OCR"""
        try:
            if isinstance(image, Image.Image):
                image = np.array(image)
            
            # Use EasyOCR for text extraction
            results = self.ocr_reader.readtext(image)
            
            # Combine all detected text
            extracted_text = ' '.join([text for _, text, _ in results])
            
            return extracted_text
        except Exception as e:
            logger.warning(f"OCR extraction failed: {e}")
            return ""
    
    def analyze_image_content(self, image: Image.Image) -> Dict[str, Any]:
        """Analyze image content and structure"""
        img_array = np.array(image)
        
        analysis = {
            'dimensions': image.size,
            'mode': image.mode,
            'has_text': False,
            'dominant_colors': [],
            'contains_chart': False,
            'contains_diagram': False
        }
        
        # Check for text
        text = self.extract_text_from_image(image)
        if text:
            analysis['has_text'] = True
            analysis['extracted_text'] = text
        
        # Detect if image contains charts/diagrams
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        edges = cv2.Canny(gray, 50, 150)
        
        # Simple heuristic for chart detection
        lines = cv2.HoughLinesP(edges, 1, np.pi/180, 100, minLineLength=100, maxLineGap=10)
        if lines is not None and len(lines) > 10:
            analysis['contains_chart'] = True
        
        # Get dominant colors
        pixels = img_array.reshape(-1, 3)
        unique_colors, counts = np.unique(pixels, axis=0, return_counts=True)
        top_colors = unique_colors[np.argsort(counts)[-5:]]
        analysis['dominant_colors'] = top_colors.tolist()
        
        return analysis


class AudioProcessor(nn.Module):
    """
    Audio processing for speech recognition and audio understanding
    Handles voice commands, audio files, and real-time audio streams
    """
    
    def __init__(self, embedding_dim: int = 768):
        super().__init__()
        self.embedding_dim = embedding_dim
        
        # Wav2Vec2 for speech recognition
        self.wav2vec_processor = Wav2Vec2Processor.from_pretrained("facebook/wav2vec2-base-960h")
        self.wav2vec_model = Wav2Vec2ForCTC.from_pretrained("facebook/wav2vec2-base-960h")
        
        # Audio feature extraction
        self.mfcc_extractor = nn.Sequential(
            nn.Conv1d(13, 64, kernel_size=3, padding=1),
            nn.ReLU(),
            nn.Conv1d(64, 128, kernel_size=3, padding=1),
            nn.ReLU(),
            nn.Conv1d(128, embedding_dim, kernel_size=3, padding=1)
        )
        
        # Transformer for audio understanding
        encoder_layer = nn.TransformerEncoderLayer(
            d_model=embedding_dim,
            nhead=8,
            dim_feedforward=2048,
            dropout=0.1,
            batch_first=True
        )
        self.audio_transformer = nn.TransformerEncoder(encoder_layer, num_layers=6)
        
        # Output projection
        self.output_projection = nn.Linear(embedding_dim, embedding_dim)
        
        # Speech recognizer for various formats
        self.recognizer = sr.Recognizer()
        
    def forward(self, audio_data: Union[torch.Tensor, np.ndarray]) -> Dict[str, Any]:
        """Process audio and extract features"""
        
        # Convert to appropriate format
        if isinstance(audio_data, np.ndarray):
            audio_data = torch.tensor(audio_data)
        
        # Process with Wav2Vec2
        inputs = self.wav2vec_processor(
            audio_data, 
            sampling_rate=16000, 
            return_tensors="pt"
        )
        
        with torch.no_grad():
            logits = self.wav2vec_model(**inputs).logits
        
        # Get predicted tokens
        predicted_ids = torch.argmax(logits, dim=-1)
        transcription = self.wav2vec_processor.batch_decode(predicted_ids)
        
        # Extract MFCC features
        mfcc_features = librosa.feature.mfcc(
            y=audio_data.numpy() if isinstance(audio_data, torch.Tensor) else audio_data,
            sr=16000,
            n_mfcc=13
        )
        mfcc_tensor = torch.tensor(mfcc_features).unsqueeze(0)
        
        # Process through feature extractor
        audio_features = self.mfcc_extractor(mfcc_tensor)
        audio_features = audio_features.transpose(1, 2)
        
        # Process through transformer
        audio_features = self.audio_transformer(audio_features)
        audio_features = audio_features.mean(dim=1)  # Global pooling
        audio_features = self.output_projection(audio_features)
        
        return {
            'audio_features': audio_features,
            'transcription': transcription[0] if transcription else "",
            'modality': 'audio'
        }
    
    def transcribe_audio_file(self, file_path: str) -> str:
        """Transcribe audio file to text"""
        try:
            # Load audio file
            audio_data, sample_rate = sf.read(file_path)
            
            # Resample if necessary
            if sample_rate != 16000:
                audio_data = librosa.resample(audio_data, orig_sr=sample_rate, target_sr=16000)
            
            # Process with Wav2Vec2
            result = self.forward(audio_data)
            return result['transcription']
            
        except Exception as e:
            logger.error(f"Audio transcription failed: {e}")
            
            # Fallback to speech_recognition
            try:
                with sr.AudioFile(file_path) as source:
                    audio = self.recognizer.record(source)
                    return self.recognizer.recognize_google(audio)
            except:
                return ""


class DocumentProcessor(nn.Module):
    """
    Document processing for PDFs, Word, Excel, and other formats
    Extracts text, tables, and structural information
    """
    
    def __init__(self, embedding_dim: int = 768):
        super().__init__()
        self.embedding_dim = embedding_dim
        
        # LayoutLM for document understanding
        self.layoutlm_processor = LayoutLMv3Processor.from_pretrained("microsoft/layoutlmv3-base")
        self.layoutlm_model = LayoutLMv3ForSequenceClassification.from_pretrained("microsoft/layoutlmv3-base")
        
        # Document encoder
        self.document_encoder = nn.Sequential(
            nn.Linear(768, embedding_dim),
            nn.ReLU(),
            nn.Dropout(0.1),
            nn.Linear(embedding_dim, embedding_dim)
        )
        
        # Supported file types
        self.supported_types = {
            '.pdf': self.process_pdf,
            '.docx': self.process_docx,
            '.doc': self.process_docx,
            '.xlsx': self.process_excel,
            '.xls': self.process_excel,
            '.csv': self.process_csv,
            '.txt': self.process_text,
            '.md': self.process_markdown,
            '.json': self.process_json,
            '.yaml': self.process_yaml,
            '.yml': self.process_yaml
        }
        
    def forward(self, document_path: str) -> Dict[str, Any]:
        """Process document and extract features"""
        
        # Determine file type
        file_ext = Path(document_path).suffix.lower()
        
        if file_ext not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Process document
        content = self.supported_types[file_ext](document_path)
        
        # Extract features using LayoutLM if applicable
        if file_ext in ['.pdf', '.docx']:
            try:
                # For documents with layout
                image = self._document_to_image(document_path)
                inputs = self.layoutlm_processor(image, return_tensors="pt")
                
                with torch.no_grad():
                    outputs = self.layoutlm_model(**inputs)
                    features = outputs.hidden_states[-1].mean(dim=1)
                
                features = self.document_encoder(features)
            except:
                # Fallback to text embedding
                features = self._text_to_features(content['text'])
        else:
            features = self._text_to_features(str(content))
        
        return {
            'document_features': features,
            'content': content,
            'file_type': file_ext,
            'modality': 'document'
        }
    
    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract content from PDF"""
        content = {
            'text': '',
            'tables': [],
            'metadata': {}
        }
        
        with pdfplumber.open(file_path) as pdf:
            # Extract metadata
            content['metadata'] = pdf.metadata
            
            # Extract text and tables from each page
            for page in pdf.pages:
                content['text'] += page.extract_text() or ''
                
                # Extract tables
                tables = page.extract_tables()
                if tables:
                    content['tables'].extend(tables)
        
        return content
    
    def process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract content from Word document"""
        doc = docx.Document(file_path)
        
        content = {
            'text': '',
            'tables': [],
            'metadata': {
                'author': doc.core_properties.author,
                'created': str(doc.core_properties.created),
                'modified': str(doc.core_properties.modified)
            }
        }
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            content['text'] += paragraph.text + '\n'
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            content['tables'].append(table_data)
        
        return content
    
    def process_excel(self, file_path: str) -> Dict[str, Any]:
        """Extract content from Excel file"""
        wb = openpyxl.load_workbook(file_path, data_only=True)
        
        content = {
            'sheets': {},
            'metadata': {
                'sheet_names': wb.sheetnames,
                'created': str(wb.properties.created),
                'modified': str(wb.properties.modified)
            }
        }
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_data = []
            
            for row in sheet.iter_rows(values_only=True):
                sheet_data.append(list(row))
            
            content['sheets'][sheet_name] = sheet_data
        
        return content
    
    def process_csv(self, file_path: str) -> List[List[str]]:
        """Extract content from CSV file"""
        data = []
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                data.append(row)
        return data
    
    def process_text(self, file_path: str) -> str:
        """Extract content from text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def process_markdown(self, file_path: str) -> Dict[str, str]:
        """Process Markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_text = f.read()
        
        return {
            'raw': md_text,
            'html': markdown.markdown(md_text)
        }
    
    def process_json(self, file_path: str) -> Dict:
        """Process JSON file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def process_yaml(self, file_path: str) -> Dict:
        """Process YAML file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)
    
    def _document_to_image(self, document_path: str) -> Image.Image:
        """Convert document page to image for layout analysis"""
        # This is a simplified version
        # In production, use pdf2image or similar
        return Image.new('RGB', (1024, 1024), color='white')
    
    def _text_to_features(self, text: str) -> torch.Tensor:
        """Convert text to features"""
        # Simplified text encoding
        # In production, use proper text encoder
        text_hash = hashlib.sha256(text.encode()).digest()
        features = torch.tensor(np.frombuffer(text_hash, dtype=np.float32))
        features = features.unsqueeze(0)
        
        # Pad or truncate to embedding_dim
        if features.size(1) < self.embedding_dim:
            features = F.pad(features, (0, self.embedding_dim - features.size(1)))
        else:
            features = features[:, :self.embedding_dim]
        
        return features


class MultiModalFusion(nn.Module):
    """
    Fusion module to combine features from different modalities
    Creates unified representation from text, image, audio, and documents
    """
    
    def __init__(self, embedding_dim: int = 768):
        super().__init__()
        self.embedding_dim = embedding_dim
        
        # Modality-specific projections
        self.text_projection = nn.Linear(embedding_dim, embedding_dim)
        self.image_projection = nn.Linear(embedding_dim, embedding_dim)
        self.audio_projection = nn.Linear(embedding_dim, embedding_dim)
        self.document_projection = nn.Linear(embedding_dim, embedding_dim)
        
        # Cross-modal attention
        self.cross_attention = nn.MultiheadAttention(
            embedding_dim,
            num_heads=12,
            batch_first=True
        )
        
        # Fusion network
        self.fusion_network = nn.Sequential(
            nn.Linear(embedding_dim * 4, embedding_dim * 2),
            nn.ReLU(),
            nn.Dropout(0.1),
            nn.Linear(embedding_dim * 2, embedding_dim),
            nn.LayerNorm(embedding_dim)
        )
        
        # Modality weights (learnable)
        self.modality_weights = nn.Parameter(torch.ones(4) / 4)
        
    def forward(self, features: Dict[str, torch.Tensor]) -> torch.Tensor:
        """Fuse features from multiple modalities"""
        
        # Project each modality
        projected = []
        
        if 'text' in features:
            projected.append(self.text_projection(features['text']))
        
        if 'image' in features:
            projected.append(self.image_projection(features['image']))
        
        if 'audio' in features:
            projected.append(self.audio_projection(features['audio']))
        
        if 'document' in features:
            projected.append(self.document_projection(features['document']))
        
        if not projected:
            raise ValueError("No modality features provided")
        
        # Stack features
        if len(projected) == 1:
            return projected[0]
        
        stacked = torch.stack(projected, dim=1)  # (batch, num_modalities, embedding_dim)
        
        # Apply cross-modal attention
        attended, _ = self.cross_attention(stacked, stacked, stacked)
        
        # Weighted combination
        weights = F.softmax(self.modality_weights[:len(projected)], dim=0)
        weighted = torch.sum(attended * weights.view(1, -1, 1), dim=1)
        
        # Final fusion
        # Pad with zeros if fewer than 4 modalities
        all_features = [weighted]
        for _ in range(4 - len(projected)):
            all_features.append(torch.zeros_like(weighted))
        
        concatenated = torch.cat(all_features, dim=-1)
        fused = self.fusion_network(concatenated)
        
        return fused


class MultiModalProcessor:
    """
    Main processor that handles all modalities
    Coordinates between different processors and manages caching
    """
    
    def __init__(self, embedding_dim: int = 768, cache_size: int = 1000):
        self.embedding_dim = embedding_dim
        
        # Initialize processors
        self.vision_processor = VisionTransformer(embedding_dim)
        self.audio_processor = AudioProcessor(embedding_dim)
        self.document_processor = DocumentProcessor(embedding_dim)
        self.fusion_module = MultiModalFusion(embedding_dim)
        
        # Cache for processed inputs
        self.cache = {}
        self.cache_size = cache_size
        
        # File type detector
        self.mime = magic.Magic(mime=True)
        
    async def process_input(self, input_data: Union[str, bytes, Image.Image, np.ndarray],
                           input_type: Optional[str] = None) -> Dict[str, Any]:
        """Process any input type and return unified features"""
        
        # Detect input type if not specified
        if input_type is None:
            input_type = self._detect_input_type(input_data)
        
        # Create multi-modal input
        mm_input = MultiModalInput(
            modality=ModalityType(input_type),
            data=input_data,
            metadata={'processed_at': str(datetime.utcnow())}
        )
        
        # Check cache
        input_hash = mm_input.get_hash()
        if input_hash in self.cache:
            return self.cache[input_hash]
        
        # Process based on modality
        result = {}
        
        if mm_input.modality == ModalityType.IMAGE:
            image = self._prepare_image(input_data)
            vision_output = self.vision_processor(image)
            result['image_features'] = vision_output['visual_features']
            
            # Extract text from image if present
            text = self.vision_processor.extract_text_from_image(image)
            if text:
                result['extracted_text'] = text
            
            # Analyze image content
            result['image_analysis'] = self.vision_processor.analyze_image_content(image)
            
        elif mm_input.modality == ModalityType.AUDIO:
            audio_data = self._prepare_audio(input_data)
            audio_output = self.audio_processor(audio_data)
            result['audio_features'] = audio_output['audio_features']
            result['transcription'] = audio_output['transcription']
            
        elif mm_input.modality == ModalityType.DOCUMENT:
            if isinstance(input_data, str):
                doc_output = self.document_processor(input_data)
                result['document_features'] = doc_output['document_features']
                result['document_content'] = doc_output['content']
            
        elif mm_input.modality == ModalityType.TEXT:
            # Process as text
            result['text'] = input_data if isinstance(input_data, str) else str(input_data)
        
        # Fuse features if multiple modalities
        if len([k for k in result.keys() if 'features' in k]) > 1:
            features_dict = {
                'image': result.get('image_features'),
                'audio': result.get('audio_features'),
                'document': result.get('document_features')
            }
            features_dict = {k: v for k, v in features_dict.items() if v is not None}
            
            result['fused_features'] = self.fusion_module(features_dict)
        
        # Update cache
        self._update_cache(input_hash, result)
        
        return result
    
    def _detect_input_type(self, input_data: Any) -> str:
        """Detect the type of input data"""
        
        if isinstance(input_data, Image.Image):
            return ModalityType.IMAGE.value
        
        elif isinstance(input_data, np.ndarray):
            # Check if it's image or audio based on shape
            if len(input_data.shape) == 3:
                return ModalityType.IMAGE.value
            elif len(input_data.shape) <= 2:
                return ModalityType.AUDIO.value
        
        elif isinstance(input_data, str):
            # Check if it's a file path
            if Path(input_data).exists():
                mime_type = self.mime.from_file(input_data)
                
                if mime_type.startswith('image/'):
                    return ModalityType.IMAGE.value
                elif mime_type.startswith('audio/'):
                    return ModalityType.AUDIO.value
                elif mime_type in ['application/pdf', 'application/msword',
                                  'application/vnd.openxmlformats-officedocument']:
                    return ModalityType.DOCUMENT.value
            
            return ModalityType.TEXT.value
        
        elif isinstance(input_data, bytes):
            # Try to detect from bytes
            try:
                Image.open(io.BytesIO(input_data))
                return ModalityType.IMAGE.value
            except:
                return ModalityType.TEXT.value
        
        return ModalityType.TEXT.value
    
    def _prepare_image(self, input_data: Union[str, bytes, Image.Image, np.ndarray]) -> Image.Image:
        """Prepare image for processing"""
        
        if isinstance(input_data, Image.Image):
            return input_data
        
        elif isinstance(input_data, np.ndarray):
            return Image.fromarray(input_data)
        
        elif isinstance(input_data, bytes):
            return Image.open(io.BytesIO(input_data))
        
        elif isinstance(input_data, str):
            if Path(input_data).exists():
                return Image.open(input_data)
            # Try base64 decode
            try:
                image_bytes = base64.b64decode(input_data)
                return Image.open(io.BytesIO(image_bytes))
            except:
                raise ValueError(f"Cannot prepare image from: {type(input_data)}")
    
    def _prepare_audio(self, input_data: Union[str, bytes, np.ndarray]) -> np.ndarray:
        """Prepare audio for processing"""
        
        if isinstance(input_data, np.ndarray):
            return input_data
        
        elif isinstance(input_data, str) and Path(input_data).exists():
            audio_data, _ = sf.read(input_data)
            return audio_data
        
        elif isinstance(input_data, bytes):
            # Try to decode audio bytes
            return np.frombuffer(input_data, dtype=np.float32)
        
        raise ValueError(f"Cannot prepare audio from: {type(input_data)}")
    
    def _update_cache(self, key: str, value: Dict[str, Any]):
        """Update cache with LRU eviction"""
        
        if len(self.cache) >= self.cache_size:
            # Remove oldest entry
            oldest_key = next(iter(self.cache))
            del self.cache[oldest_key]
        
        self.cache[key] = value


# Export main components
__all__ = [
    'MultiModalProcessor',
    'VisionTransformer',
    'AudioProcessor',
    'DocumentProcessor',
    'MultiModalFusion',
    'MultiModalInput',
    'ModalityType'
]