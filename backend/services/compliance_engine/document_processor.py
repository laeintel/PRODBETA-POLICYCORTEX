"""
Document Processing Pipeline for Policy Compliance Engine
Handles document upload, storage, and processing with Azure Blob Storage and Azure Functions
"""

import asyncio
import hashlib
import io
import json
import mimetypes
from datetime import datetime
from datetime import timedelta
from enum import Enum
from pathlib import Path
from typing import Any
from typing import BinaryIO
from typing import Dict
from typing import List
from typing import Optional

import docx  # python-docx for Word processing
import fitz  # PyMuPDF for PDF processing
import pandas as pd
import structlog
from azure.functions import HttpRequest
from azure.functions import HttpResponse
from azure.identity.aio import DefaultAzureCredential
from azure.storage.blob import BlobSasPermissions
from azure.storage.blob import generate_blob_sas
from azure.storage.blob.aio import BlobClient
from azure.storage.blob.aio import BlobServiceClient
from pydantic import BaseModel
from pydantic import Field
from pydantic import validator

logger = structlog.get_logger(__name__)


class DocumentType(str, Enum):
    """Supported document types for processing"""

    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    CSV = "csv"
    JSON = "json"
    XLSX = "xlsx"
    HTML = "html"
    MARKDOWN = "md"


class DocumentStatus(str, Enum):
    """Document processing status"""

    UPLOADED = "uploaded"
    PROCESSING = "processing"
    ANALYZED = "analyzed"
    FAILED = "failed"
    ARCHIVED = "archived"


class ProcessedDocument(BaseModel):
    """Processed document metadata and content"""

    document_id: str
    tenant_id: str
    filename: str
    document_type: DocumentType
    status: DocumentStatus
    size_bytes: int
    content_hash: str
    storage_url: str
    extracted_text: Optional[str] = None
    metadata: Dict[str, Any] = Field(default_factory=dict)
    processing_errors: List[str] = Field(default_factory=list)
    uploaded_at: datetime
    processed_at: Optional[datetime] = None

    class Config:
        use_enum_values = True


class DocumentProcessor:
    """
    Handles document upload, storage, and processing pipeline
    Integrates with Azure Blob Storage and Azure Functions for scalable processing
    """

    def __init__(
        self,
        storage_account_name: str,
        container_name: str = "policy-documents",
        credential: Optional[DefaultAzureCredential] = None,
    ):
        self.storage_account_name = storage_account_name
        self.container_name = container_name
        self.credential = credential or DefaultAzureCredential()
        self.blob_service_client = None
        self.supported_types = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".txt": DocumentType.TXT,
            ".csv": DocumentType.CSV,
            ".json": DocumentType.JSON,
            ".xlsx": DocumentType.XLSX,
            ".html": DocumentType.HTML,
            ".md": DocumentType.MARKDOWN,
        }

    async def initialize(self):
        """Initialize blob service client"""
        if not self.blob_service_client:
            account_url = f"https://{self.storage_account_name}.blob.core.windows.net"
            self.blob_service_client = BlobServiceClient(
                account_url=account_url, credential=self.credential
            )

            # Ensure container exists
            container_client = self.blob_service_client.get_container_client(self.container_name)

            try:
                await container_client.create_container()
                logger.info(f"Created container: {self.container_name}")
            except Exception as e:
                if "ContainerAlreadyExists" not in str(e):
                    logger.error(f"Error creating container: {e}")
                    raise

    async def upload_document(
        self,
        file_content: BinaryIO,
        filename: str,
        tenant_id: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> ProcessedDocument:
        """
        Upload document to Azure Blob Storage and initiate processing

        Args:
            file_content: Binary file content
            filename: Original filename
            tenant_id: Tenant identifier for multi-tenancy
            metadata: Additional metadata to store with document

        Returns:
            ProcessedDocument with initial metadata
        """
        await self.initialize()

        # Detect document type
        file_extension = Path(filename).suffix.lower()
        if file_extension not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_extension}")

        document_type = self.supported_types[file_extension]

        # Generate document ID and blob name
        content_bytes = file_content.read()
        file_content.seek(0)  # Reset for further processing

        content_hash = hashlib.sha256(content_bytes).hexdigest()
        document_id = (
            f"{tenant_id}/{datetime.utcnow().strftime('%Y%m%d')}/{content_hash[:8]}_{filename}"
        )

        # Upload to blob storage
        blob_client = self.blob_service_client.get_blob_client(
            container=self.container_name, blob=document_id
        )

        upload_metadata = {
            "tenant_id": tenant_id,
            "original_filename": filename,
            "document_type": document_type.value,
            "content_hash": content_hash,
            "uploaded_at": datetime.utcnow().isoformat(),
        }

        if metadata:
            upload_metadata.update(metadata)

        try:
            await blob_client.upload_blob(file_content, overwrite=True, metadata=upload_metadata)

            logger.info(f"Uploaded document: {document_id}")

        except Exception as e:
            logger.error(f"Failed to upload document: {e}")
            raise

        # Create document record
        storage_url = self._generate_sas_url(document_id)

        processed_doc = ProcessedDocument(
            document_id=document_id,
            tenant_id=tenant_id,
            filename=filename,
            document_type=document_type,
            status=DocumentStatus.UPLOADED,
            size_bytes=len(content_bytes),
            content_hash=content_hash,
            storage_url=storage_url,
            metadata=metadata or {},
            uploaded_at=datetime.utcnow(),
        )

        # Trigger async processing
        asyncio.create_task(self._process_document_async(processed_doc, content_bytes))

        return processed_doc

    def _generate_sas_url(self, blob_name: str, expiry_hours: int = 24) -> str:
        """Generate SAS URL for secure document access"""
        sas_token = generate_blob_sas(
            account_name=self.storage_account_name,
            container_name=self.container_name,
            blob_name=blob_name,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.utcnow() + timedelta(hours=expiry_hours),
        )

        return f"https://{self.storage_account_name}.blob.core.windows.net/{self.container_name}/{blob_name}?{sas_token}"

    async def _process_document_async(self, document: ProcessedDocument, content: bytes):
        """Process document asynchronously"""
        try:
            document.status = DocumentStatus.PROCESSING

            # Extract text based on document type
            extracted_text = await self._extract_text(content, document.document_type)

            document.extracted_text = extracted_text
            document.status = DocumentStatus.ANALYZED
            document.processed_at = datetime.utcnow()

            # Update metadata in blob storage
            await self._update_document_metadata(document)

            logger.info(f"Processed document: {document.document_id}")

        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            document.status = DocumentStatus.FAILED
            document.processing_errors.append(str(e))

    async def _extract_text(self, content: bytes, document_type: DocumentType) -> str:
        """Extract text content from various document formats"""

        if document_type == DocumentType.PDF:
            return self._extract_pdf_text(content)

        elif document_type == DocumentType.DOCX:
            return self._extract_docx_text(content)

        elif document_type == DocumentType.TXT:
            return content.decode("utf-8", errors="ignore")

        elif document_type == DocumentType.CSV:
            df = pd.read_csv(io.BytesIO(content))
            return df.to_string()

        elif document_type == DocumentType.JSON:
            data = json.loads(content)
            return json.dumps(data, indent=2)

        elif document_type == DocumentType.XLSX:
            df = pd.read_excel(io.BytesIO(content))
            return df.to_string()

        elif document_type in [DocumentType.HTML, DocumentType.MARKDOWN]:
            return content.decode("utf-8", errors="ignore")

        else:
            raise ValueError(f"Unsupported document type: {document_type}")

    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF using PyMuPDF"""
        text_content = []

        try:
            pdf_document = fitz.open(stream=content, filetype="pdf")

            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text = page.get_text()
                text_content.append(f"--- Page {page_num + 1} ---\n{text}")

            pdf_document.close()

        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise

        return "\n\n".join(text_content)

    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from Word document"""
        text_content = []

        try:
            doc = docx.Document(io.BytesIO(content))

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text_content.append(" | ".join(row_text))

        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise

        return "\n\n".join(text_content)

    async def _update_document_metadata(self, document: ProcessedDocument):
        """Update document metadata in blob storage"""
        blob_client = self.blob_service_client.get_blob_client(
            container=self.container_name, blob=document.document_id
        )

        metadata = {
            "status": document.status.value,
            "processed_at": document.processed_at.isoformat() if document.processed_at else "",
            "text_length": str(len(document.extracted_text)) if document.extracted_text else "0",
            "has_errors": "true" if document.processing_errors else "false",
        }

        await blob_client.set_blob_metadata(metadata)

    async def get_document(self, document_id: str) -> ProcessedDocument:
        """Retrieve document metadata and content"""
        await self.initialize()

        blob_client = self.blob_service_client.get_blob_client(
            container=self.container_name, blob=document_id
        )

        # Get blob properties and metadata
        properties = await blob_client.get_blob_properties()

        # Download content if needed
        download_stream = await blob_client.download_blob()
        content = await download_stream.readall()

        # Extract text
        document_type = DocumentType(properties.metadata.get("document_type"))
        extracted_text = await self._extract_text(content, document_type)

        return ProcessedDocument(
            document_id=document_id,
            tenant_id=properties.metadata.get("tenant_id"),
            filename=properties.metadata.get("original_filename"),
            document_type=document_type,
            status=DocumentStatus(properties.metadata.get("status", "uploaded")),
            size_bytes=properties.size,
            content_hash=properties.metadata.get("content_hash"),
            storage_url=self._generate_sas_url(document_id),
            extracted_text=extracted_text,
            metadata=properties.metadata,
            uploaded_at=properties.last_modified,
            processed_at=(
                datetime.fromisoformat(properties.metadata.get("processed_at"))
                if properties.metadata.get("processed_at")
                else None
            ),
        )

    async def list_documents(
        self, tenant_id: str, status: Optional[DocumentStatus] = None, limit: int = 100
    ) -> List[ProcessedDocument]:
        """List documents for a tenant"""
        await self.initialize()

        documents = []
        container_client = self.blob_service_client.get_container_client(self.container_name)

        # List blobs with prefix filter for tenant
        prefix = f"{tenant_id}/"

        async for blob in container_client.list_blobs(
            name_starts_with=prefix, include=["metadata"], results_per_page=limit
        ):
            if status and blob.metadata.get("status") != status.value:
                continue

            doc = ProcessedDocument(
                document_id=blob.name,
                tenant_id=tenant_id,
                filename=blob.metadata.get("original_filename", blob.name),
                document_type=DocumentType(blob.metadata.get("document_type", "txt")),
                status=DocumentStatus(blob.metadata.get("status", "uploaded")),
                size_bytes=blob.size,
                content_hash=blob.metadata.get("content_hash", ""),
                storage_url=self._generate_sas_url(blob.name),
                metadata=blob.metadata,
                uploaded_at=blob.last_modified,
            )

            documents.append(doc)

            if len(documents) >= limit:
                break

        return documents

    async def delete_document(self, document_id: str) -> bool:
        """Delete a document from storage"""
        await self.initialize()

        blob_client = self.blob_service_client.get_blob_client(
            container=self.container_name, blob=document_id
        )

        try:
            await blob_client.delete_blob()
            logger.info(f"Deleted document: {document_id}")
            return True

        except Exception as e:
            logger.error(f"Failed to delete document: {e}")
            return False


# Azure Function for document processing
async def process_document_function(req: HttpRequest) -> HttpResponse:
    """
    Azure Function endpoint for document processing
    Triggered via HTTP request with document details
    """
    try:
        req_body = req.get_json()
        document_id = req_body.get("document_id")
        storage_account = req_body.get("storage_account")
        container = req_body.get("container", "policy-documents")

        if not document_id or not storage_account:
            return HttpResponse(
                json.dumps({"error": "Missing required parameters"}),
                status_code=400,
                mimetype="application/json",
            )

        # Initialize processor
        processor = DocumentProcessor(
            storage_account_name=storage_account, container_name=container
        )

        # Process document
        document = await processor.get_document(document_id)

        return HttpResponse(
            json.dumps(
                {
                    "document_id": document.document_id,
                    "status": document.status.value,
                    "extracted_text_length": (
                        len(document.extracted_text) if document.extracted_text else 0
                    ),
                    "processed_at": (
                        document.processed_at.isoformat() if document.processed_at else None
                    ),
                }
            ),
            status_code=200,
            mimetype="application/json",
        )

    except Exception as e:
        logger.error(f"Function processing error: {e}")
        return HttpResponse(
            json.dumps({"error": str(e)}), status_code=500, mimetype="application/json"
        )
